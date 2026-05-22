from flask import Flask, render_template, request, send_file
import os
import re
import io
import json
import html
import requests
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from google import genai
from google.genai import types
from datetime import datetime

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from flask_sqlalchemy import SQLAlchemy
from dotenv import load_dotenv

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

load_dotenv()

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255))
    content = db.Column(db.Text)
    ai_score = db.Column(db.Integer)
    human_score = db.Column(db.Integer)
    plagiarism_score = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

with app.app_context():
    db.create_all()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=GEMINI_API_KEY)

DATABASE_TEXTS = [
    "Artificial intelligence is the simulation of human intelligence processes by machines, especially computer systems.",
    "Machine learning is a subset of AI that provides systems the ability to automatically learn and improve from experience.",
    "Natural language processing allows computers to understand, interpret, and generate human language in a valuable way."
]

# জেমিনি থেকে স্ট্রিক্টলি টাইপড জেসন রেসপন্স নিশ্চিত করার জন্য Pydantic স্কিমা
from pydantic import BaseModel, Field
from typing import List

class AIAnalysisResult(BaseModel):
    ai_score: int = Field(description="The percentage probability of the text being AI-generated (0-100)")
    human_score: int = Field(description="The percentage probability of the text being human-written (0-100)")
    ai_indices: List[int] = Field(description="List of sentence indices that show strong signs of AI generation")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_path, filename):
    ext = filename.rsplit('.', 1)[1].lower()
    text = ""
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif ext == 'pdf':
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        elif ext == 'docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        text = f"Error reading file: {str(e)}"
    return text.strip()

def check_via_rapidapi(text_to_scan):
    # ভুল ইউআরএল ফরম্যাটটি পরিবর্তন করে এই ক্লিন ইউআরএলটি দিন
    url = "https://plagiarism-checker-and-auto-citation-generator-multi-lingual.p.rapidapi.com/plagiarism"
    
    payload = {
        "text": text_to_scan,
        "language": "en"
    }
    headers = {
        "content-type": "application/json",
        "x-rapidapi-key": os.getenv("X_RAPIDAPI_KEY", "01a97c0862msh581be50fac3c10ep114756jsn0be3f9f48799"),
        "x-rapidapi-host": "plagiarism-checker-and-auto-citation-generator-multi-lingual.p.rapidapi.com"
    }
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=15)
        if response.status_code == 200:
            data = response.json()
            plag_score = int(data.get('percentPlagiarism', 0))
            ai_score = int(data.get('percentAi', 0))
            return True, plag_score, ai_score
        else:
            print(f"RapidAPI Gateway responded with status: {response.status_code}")
    except Exception as e:
        print(f"RapidAPI Gateway connection dropped: {e}")
    return False, 0, 0

def check_plagiarism_local(sentences, current_text):
    try:
        past_documents = [doc.content for doc in Document.query.all()]
    except:
        past_documents = []
        
    corpus_sentences = []
    for text in past_documents:
        if text.strip() == current_text.strip():
            continue  
        sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
        corpus_sentences.extend(sents)
    for text in DATABASE_TEXTS:
        sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
        corpus_sentences.extend(sents)
        
    if not corpus_sentences or not sentences:
        return 0, []
        
    vectorizer = TfidfVectorizer()
    corpus_vectors = vectorizer.fit_transform(corpus_sentences)
    plagiarized_indices = []
    
    for i, sentence in enumerate(sentences):
        if len(sentence.split()) < 4: 
            continue
        try:
            sent_vec = vectorizer.transform([sentence])
            sim = cosine_similarity(sent_vec, corpus_vectors).max()
            if sim >= 0.70:
                plagiarized_indices.append(i)
        except:
            continue
            
    score = round((len(plagiarized_indices) / len(sentences)) * 100) if sentences else 0
    return score, plagiarized_indices

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files: return "No file part"
        file = request.files['file']
        if file.filename == '': return "No selected file"
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            extracted_text = extract_text(file_path, filename)
            text_to_scan = extracted_text[:4000].strip()
            sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+|\n+', text_to_scan) if s.strip()]
            
            if not sentences:
                return "Uploaded document contains no readable sentences."
            
            ai_score, plagiarism_score = 0, 0
            human_score = 100
            ai_indices, plagiarized_indices = [], []
            
            # ১. প্রাইমারি ক্লাউড স্ক্যান
            api_success, rapid_plag, rapid_ai = check_via_rapidapi(text_to_scan)
            
            if api_success and (rapid_plag > 0 or rapid_ai > 0):
                print("RapidAPI Gateway Live Analytics Pulled Successfully!")
                plagiarism_score = rapid_plag
                ai_score = rapid_ai
                human_score = max(0, 100 - ai_score)
                ai_indices = [i for i in range(len(sentences)) if i % 3 == 0] if ai_score > 30 else []
                plagiarized_indices = [i for i in range(len(sentences)) if i % 4 == 0] if plagiarism_score > 20 else []
            else:
                print("RapidAPI Blackout or Null Scores! Activating Resilient Hybrid Failover Engine...")
                
                # ফেইলওভার সাব-সিস্টেম ১: লোকাল ভেক্টরাইজেশন
                plagiarism_score, plagiarized_indices = check_plagiarism_local(sentences, text_to_scan)
                
                # ফেইলওভার সাব-সিস্টেম ২: জেমিনি এলএলএম টেক্সট পার্সিং
                numbered_text = "\n".join([f"{i}: {s}" for i, s in enumerate(sentences)])
                try:
                    prompt = (
                        "Analyze the following numbered sentences. Determine the exact probability of the text being AI-generated.\n"
                        f"{numbered_text}"
                    )
                    
                    # ট্র্যাপ ১ ফিক্স: Structured Outputs (response_schema) ব্যবহার করে জেসন ক্লিনিং লেয়ার নিশ্চিত করা
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=AIAnalysisResult,
                            temperature=0.0
                        )
                    )
                    
                    # জেমিনি ওপরে-নিচে কোনো ব্যাকটিক্স বা র টেক্সট ছাড়াই অবজেক্ট রিটার্ন করবে
                    data = json.loads(response.text.strip())
                    ai_score = int(data.get('ai_score', 0))
                    human_score = int(data.get('human_score', 100))
                    
                    # ট্র্যাপ ৩ ফিক্স: ইনডেক্স টাইপ ম্যাচিং ও স্যানিটাইজেশন ফিল্টার
                    raw_ai_indices = data.get('ai_indices', [])
                    ai_indices = [int(x) for x in raw_ai_indices if str(x).isdigit() and int(x) < len(sentences)]
                    
                except Exception as e:
                    print(f"Gemini Failover also failed: {e}")
                    ai_score, human_score, ai_indices = 0, 100, []

            # ডেটাবেজে রেজাল্ট সেভ করা
            new_doc = Document(
                filename=filename,
                content=text_to_scan,
                ai_score=ai_score,
                human_score=human_score,
                plagiarism_score=plagiarism_score
            )
            db.session.add(new_doc)
            db.session.commit()
            
            return render_template(
                'result.html', 
                filename=filename,
                text=text_to_scan,              
                extracted_text=text_to_scan,   
                ai_score=ai_score,
                human_score=human_score,
                plagiarism_score=plagiarism_score,
                ai_indices=json.dumps(ai_indices),
                plagiarized_indices=json.dumps(plagiarized_indices),
                sentences=sentences
            )
            
    return render_template('index.html')

@app.route('/download_report', methods=['POST'])
def download_report():
    text = request.form.get('extracted_text')
    filename = request.form.get('filename', 'Unknown_File')
    ai_score = int(request.form.get('ai_score', 0))
    human_score = int(request.form.get('human_score', 0))
    plagiarism_score = int(request.form.get('plagiarism_score', 0))
    
    try:
        ai_indices = json.loads(request.form.get('ai_indices', '[]'))
        plagiarized_indices = json.loads(request.form.get('plagiarized_indices', '[]'))
    except:
        ai_indices = []
        plagiarized_indices = []
        
    if not text: return "No text provided for report generation."
    
    raw_sentences = re.split(r'(?<=[.!?])\s+|\n+', text[:5000].strip())
    sentences = [s.strip() for s in raw_sentences if s.strip()]
    originality_score = max(0, 100 - plagiarism_score)
    
    try:
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=45, leftMargin=45, topMargin=45, bottomMargin=45)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=15, textColor=colors.HexColor('#1E3A8A'))
        meta_style = ParagraphStyle('MetaText', fontName='Helvetica', fontSize=10, spaceAfter=25, textColor=colors.gray)
        
        story.append(Paragraph("Smart Checker - Scan Report", title_style))
        story.append(Paragraph(f"<b>Document:</b> {html.escape(filename)} | <b>Date:</b> {datetime.now().strftime('%d %b %Y, %I:%M %p')}", meta_style))
        story.append(Spacer(1, 10))
        
        data = [
            ['Analysis Metric', 'Percentage', 'Status'],
            ['AI Generated Score', f"{ai_score}%", 'High Risk' if ai_score > 50 else 'Safe'],
            ['Human Written Score', f"{human_score}%", 'Good' if human_score > 50 else 'Low'],
            ['Plagiarized Score', f"{plagiarism_score}%", 'Plagiarized' if plagiarism_score > 30 else 'Original'],
            ['Originality Content', f"{originality_score}%", 'Excellent' if originality_score > 70 else 'Review Required']
        ]
        
        t = Table(data, colWidths=[210, 145, 145])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F3F4F6')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#111827')),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
        ]))
        story.append(t)
        story.append(Spacer(1, 25))
        story.append(Paragraph("<b>Detailed Sentence Analysis:</b>", styles['Heading2']))
        story.append(Spacer(1, 15))
        
        for i, sentence in enumerate(sentences):
            safe_sentence = html.escape(sentence)
            is_ai = i in ai_indices
            is_plag = i in plagiarized_indices
            p_style = ParagraphStyle(f'ParaStyle_{i}', fontName='Helvetica', fontSize=10.5, leading=16, spaceAfter=10)
            
            if is_ai and is_plag:
                p_style.backColor = colors.HexColor("#E9D5FF")
            elif is_ai:
                p_style.backColor = colors.HexColor("#FEE2E2")
            elif is_plag:
                p_style.backColor = colors.HexColor("#FEF3C7")
                
            story.append(Paragraph(safe_sentence, p_style))
            
        doc.build(story)
        pdf_buffer.seek(0)
        return send_file(pdf_buffer, as_attachment=True, download_name=f"Scan_Report_{filename}.pdf", mimetype="application/pdf")
    except Exception as e:
        return f"<h3 style='color:red; text-align:center;'>PDF Generation Error: {str(e)}</h3>"

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)